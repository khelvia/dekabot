import os
import threading
from flask import Flask
import pandas as pd
import yfinance as yf
from docx import Document
from telegram import Update
from telegram.ext import (
    ApplicationBuilder,
    MessageHandler,
    CommandHandler,
    ContextTypes,
    filters,
)
from google import genai

# ==============================
# ENV VARIABLES
# ==============================

TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not TELEGRAM_TOKEN or not GEMINI_API_KEY:
    raise ValueError("Missing TELEGRAM_TOKEN or GEMINI_API_KEY")

# Gemini 2.5 Client
client = genai.Client(api_key=GEMINI_API_KEY)

# ==============================
# FLASK SERVER (FOR RENDER PORT)
# ==============================

def run_web():
    web_app = Flask(__name__)

    @web_app.route("/")
    def home():
        return "Telegram AI Bot is running."

    port = int(os.environ.get("PORT", 10000))
    web_app.run(host="0.0.0.0", port=port)


# ==============================
# GEMINI HELPER FUNCTION
# ==============================

def generate_ai_response(prompt: str) -> str:
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
    )
    return response.text


# ==============================
# TEXT CHAT
# ==============================

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        user_text = update.message.text
        reply = generate_ai_response(user_text)
        await update.message.reply_text(reply)
    except Exception as e:
        await update.message.reply_text(f"Error: {str(e)}")


# ==============================
# WORD REWRITE
# ==============================

async def rewrite_doc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        file = await update.message.document.get_file()
        input_path = "input.docx"
        output_path = "rewritten.docx"

        await file.download_to_drive(input_path)

        doc = Document(input_path)
        text = "\n".join([p.text for p in doc.paragraphs])

        if not text.strip():
            await update.message.reply_text("Document is empty.")
            return

        rewritten = generate_ai_response(
            f"Rewrite this professionally while preserving meaning:\n\n{text}"
        )

        new_doc = Document()
        new_doc.add_paragraph(rewritten)
        new_doc.save(output_path)

        await update.message.reply_document(open(output_path, "rb"))

    except Exception as e:
        await update.message.reply_text(f"Word processing error: {str(e)}")


# ==============================
# EXCEL ANALYSIS
# ==============================

async def handle_excel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        file = await update.message.document.get_file()
        input_path = "input.xlsx"
        output_path = "analysis.xlsx"

        await file.download_to_drive(input_path)

        df = pd.read_excel(input_path)
        summary = df.describe(include="all")

        writer = pd.ExcelWriter(output_path, engine="openpyxl")
        df.to_excel(writer, sheet_name="Original", index=False)
        summary.to_excel(writer, sheet_name="Summary")
        writer.close()

        await update.message.reply_document(open(output_path, "rb"))

    except Exception as e:
        await update.message.reply_text(f"Excel processing error: {str(e)}")


# ==============================
# STOCK ANALYSIS
# ==============================

async def stock(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        if len(context.args) == 0:
            await update.message.reply_text("Usage: /stock TCS.NS")
            return

        ticker = context.args[0]
        data = yf.download(ticker, period="6mo", progress=False)

        if data.empty:
            await update.message.reply_text("Invalid or unavailable ticker.")
            return

        returns = data["Close"].pct_change().mean()
        volatility = data["Close"].pct_change().std()

        prompt = f"""
        Stock: {ticker}
        Average Daily Return: {returns}
        Volatility: {volatility}

        Provide interpretation, risk assessment, and outlook.
        """

        analysis = generate_ai_response(prompt)

        await update.message.reply_text(analysis)

    except Exception as e:
        await update.message.reply_text(f"Stock analysis error: {str(e)}")


# ==============================
# MAIN FUNCTION
# ==============================

def main():
    app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()

    # Commands
    app.add_handler(CommandHandler("stock", stock))

    # Excel upload
    app.add_handler(
        MessageHandler(
            filters.Document.MimeType(
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ),
            handle_excel,
        )
    )

    # Word upload
    app.add_handler(
        MessageHandler(
            filters.Document.MimeType(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            rewrite_doc,
        )
    )

    # Text messages
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    print("Bot started successfully.")

    # Start Flask (for Render port binding)
    threading.Thread(target=run_web).start()

    # Start Telegram polling
    app.run_polling()


if __name__ == "__main__":
    main()